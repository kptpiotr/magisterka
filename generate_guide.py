from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.shared import Cm, Inches


class GuideGenerator:
    def __init__(self, template_file_name='template.docx'):
        # Load the template document
        self.doc = Document(template_file_name)

    def replace(self, to_replace: dict):
        # ['{NAZWA_FIRMY}', '{NR}', '{NR_CZESCI}', '{NR_CZESCI}', '{NAZWA_CZESCI}', '{NR_ZAMOW}', '{NR_RYS}', '{NR_ZLECENIA}', '{NR_PL}', '{AMOUNT}', '{UWAGA}', '{TERMIN_DOSTAWY}', '{DATA_EMISJI}', '{MATERIAL }', '{J}', '{NORMA}', '{ODPAD}', '{WYMIARY}', '{NR_PARTII}', '{NR_WYTOPU}', '{NR_PZ}', '{GATUNEK}', '{NORMA_MAT}', '{DOD_MAT}']
        # First edit the text
        # Iterate through tables in the document
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key in to_replace:
                        if '{' + key + '}' in cell.text:
                            # print(f"Replacing {key} to {to_replace[key]}")
                            cell.text = cell.text.replace('{' + key + '}', to_replace[key])

    def add_header_table(self):

        # Add a table with 1 row and 9 columns
        table = self.doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'

        # table.allow_autofit = True
        #
        # for col in table.columns:
        #     col.width = Inches(0.83)
        # table.columns[0].width = Inches(0.83)
        # table.columns[1].width = Inches(0.83)
        # table.columns[2].width = Inches(0.83)
        # table.columns[3].width = Inches(0.83)
        # table.columns[4].width = Inches(0.83)
        # table.columns[5].width = Inches(0.83)
        # table.columns[0].width = Inches(0.83)
        # table.columns[0].width = Inches(0.83)
        # table.columns[0].width = Inches(0.83)

        # Populate the first row with the specified texts
        row = table.rows[0]
        row.cells[0].text = "Oper/czyn"
        row.cells[1].text = "Sym oper."
        row.cells[2].text = "Nr zn. wyokn."
        row.cells[3].text = "ODB 1 SZT. KJ"
        row.cells[4].text = "MISTRZ"
        row.cells[5].text = "DOBRE"
        row.cells[6].text = "ZŁE"
        row.cells[7].text = "KJ"
        row.cells[8].text = "UWAGI"

        # break_paragraph = self.doc.add_paragraph()
        # break_run = break_paragraph.add_run("\n")

    def add_process_tables(self, df):
        df.fillna(" ", inplace=True)
        # Create new tables and set the cell text and column widths
        for index, row in df.iterrows():
            new_table = self.doc.add_table(rows=1, cols=9)

            # Apply a table style that includes borders
            new_table.style = 'Table Grid'

            # Set the cell text for the new table
            # for table_row in new_table.rows:
            #     for idx, table_cell in enumerate(table_row.cells):
            #         print(row)
            #         table_cell.text = row[idx]
            #
            #         # Set cell alignment
            #         table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            # print(row)

            for table_row in new_table.rows:
                try:
                    table_row.cells[0].text = row[0]
                except TypeError:
                    pass

                try:
                    table_row.cells[1].text = row[1]
                except TypeError:
                    pass

                for idx, table_cell in enumerate(table_row.cells):
                    # Set cell alignment
                    table_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Add the text
            self.doc.add_paragraph('NAZWA: ' + row[2])
            self.doc.add_paragraph('CZAS: ' + row[4])
            self.doc.add_paragraph(row[3])

    def add_ending(self):
        table = self.doc.add_table(rows=3, cols=4)
        table.style = 'Table Grid'

        # Populate the first row with the specified texts
        row0 = table.rows[0]
        row0.cells[1].text = "Opracował"
        row0.cells[2].text = "Zatwierdził"
        row0.cells[3].text = "KJ - Zatwierdził"

        row1 = table.rows[1]
        row1.cells[0].text = "Podpis"

        row2 = table.rows[2]
        row2.cells[0].text = "Data"

    def add_tools_page(self, df, company_name, guide_number, part_number, date):
        # Add a new page
        self.doc.add_page_break()

        # Add a table with 2 rows and 3 columns to the second page
        header_table = self.doc.add_table(rows=2, cols=3)
        header_table.style = 'Table Grid'

        header_table.rows[0].cells[0].text = company_name
        header_table.rows[0].cells[1].text = "SPECYFIKACJA NARZĘDZI I OPRZYRZĄDOWANIA"
        header_table.rows[1].cells[0].text = "NR PRZEWODNIKA\n" + guide_number
        header_table.rows[1].cells[1].text = "NR CZĘŚCI\n" + part_number
        header_table.rows[1].cells[2].text = "DATA\n" + date

        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'

        header_row = table.rows[0]
        header_row.cells[0].text = "NR OPERACJI"
        header_row.cells[1].text = "NAZWA NARZĘDZIA"
        header_row.cells[2].text = "NR NARZĘDZIA"
        header_row.cells[3].text = "OPIS NARZĘDZIA"

        for index, row in df.iterrows():
            table_row = table.add_row()
            table_row.cells[0].text = str(row[0])
            table_row.cells[2].text = str(row[1])
            table_row.cells[1].text = str(row[2])
            table_row.cells[3].text = str(row[3])

    def add_materials_page(self, df, company_name, guide_number, part_number, date):
        # Add a new page
        self.doc.add_page_break()

        # Add a table with 2 rows and 3 columns to the second page
        header_table = self.doc.add_table(rows=2, cols=3)
        header_table.style = 'Table Grid'

        header_table.rows[0].cells[0].text = company_name
        header_table.rows[0].cells[1].text = "SPECYFIKACJA PÓŁFABRYKATÓW I MATERIAŁÓW DO MONTAŻU"
        header_table.rows[1].cells[0].text = "NR PRZEWODNIKA\n" + guide_number
        header_table.rows[1].cells[1].text = "NR CZĘŚCI\n" + part_number
        header_table.rows[1].cells[2].text = "DATA\n" + date

        table = self.doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'

        header_row = table.rows[0]
        header_row.cells[0].text = "NR OPERACJI"
        header_row.cells[1].text = "NR MATERIAŁU"
        header_row.cells[2].text = "NAZWA MATERIAŁU"
        header_row.cells[3].text = "NR PARTII"
        header_row.cells[4].text = "JED M."
        header_row.cells[5].text = "NORMA/1 SZT"
        header_row.cells[6].text = "NORMA PRZEWODNIK"
        header_row.cells[7].text = "WYDANO"

        for index, row in df.iterrows():
            table_row = table.add_row()
            table_row.cells[0].text = str(row[0])
            table_row.cells[1].text = str(row[1])
            table_row.cells[2].text = str(row[2])
            table_row.cells[3].text = str(row[3])
            table_row.cells[4].text = str(row[4])
            table_row.cells[5].text = str(row[5])
            table_row.cells[6].text = str(row[6])
            table_row.cells[7].text = str(row[7])

        self.doc.add_paragraph("WYDAŁ \nDATA PODPIS\n\n\n")
        self.doc.add_paragraph("POBRAŁ \nDATA PODPIS")

    def save_file(self, output_file_name='output.docx'):
        # Save the modified document
        self.doc.save(output_file_name)


if __name__ == "__main__":
    generator = GuideGenerator()
    generator.replace({'{NAZWA_FIRMY}': "WZK"})
    generator.add_header_table()
    generator.add_process_tables(1)
    generator.save_file()

    # x = {'{NAZWA_FIRMY}': "WZK"}
    # print(x[1])

